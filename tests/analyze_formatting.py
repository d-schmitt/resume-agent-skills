#!/usr/bin/env python3
"""Analyze DOCX formatting and output structured JSON.

Usage: python analyze_formatting.py <path-to-docx> [--json]
"""

import json
import sys
from docx import Document
from docx.oxml.ns import qn


def emu_to_inches(emu_val):
    if emu_val is None:
        return None
    return round(emu_val / 914400, 3)


def color_str(color):
    if color is None:
        return None
    try:
        if color.rgb is not None:
            return str(color.rgb)
    except Exception:
        pass
    try:
        if color.theme_color is not None:
            return f"theme:{color.theme_color}"
    except Exception:
        pass
    return None


def analyze(docx_path):
    doc = Document(docx_path)
    result = {}

    # Page setup
    sections = []
    for i, s in enumerate(doc.sections):
        sections.append({
            "index": i,
            "page_width_in": emu_to_inches(s.page_width),
            "page_height_in": emu_to_inches(s.page_height),
            "margin_top_in": emu_to_inches(s.top_margin),
            "margin_bottom_in": emu_to_inches(s.bottom_margin),
            "margin_left_in": emu_to_inches(s.left_margin),
            "margin_right_in": emu_to_inches(s.right_margin),
        })
    result["page_setup"] = sections

    # Doc defaults
    defaults = {}
    dd = doc.styles.element.find(qn("w:docDefaults"))
    if dd is not None:
        rprd = dd.find(qn("w:rPrDefault"))
        if rprd is not None:
            rpr = rprd.find(qn("w:rPr"))
            if rpr is not None:
                rf = rpr.find(qn("w:rFonts"))
                if rf is not None:
                    for attr in ["ascii", "hAnsi", "eastAsia", "cs"]:
                        v = rf.get(qn(f"w:{attr}"))
                        if v:
                            defaults[f"font_{attr}"] = v
                    for attr in ["asciiTheme", "hAnsiTheme", "eastAsiaTheme", "cstheme"]:
                        v = rf.get(qn(f"w:{attr}"))
                        if v:
                            defaults[f"theme_{attr}"] = v
                sz = rpr.find(qn("w:sz"))
                if sz is not None:
                    defaults["font_size_pt"] = int(sz.get(qn("w:val"))) / 2
    result["doc_defaults"] = defaults

    # Styles actually used
    para_styles_used = set()
    for p in doc.paragraphs:
        if p.style:
            para_styles_used.add(p.style.name)
    result["para_styles_used"] = sorted(para_styles_used)

    # Style definitions (only used ones)
    style_defs = {}
    for style in doc.styles:
        if style.name not in para_styles_used:
            continue
        sd = {"base": style.base_style.name if style.base_style else None}
        try:
            f = style.font
            if f:
                if f.name:
                    sd["font_name"] = f.name
                if f.size:
                    sd["font_size_pt"] = f.size.pt
                if f.bold is not None:
                    sd["bold"] = f.bold
                if f.italic is not None:
                    sd["italic"] = f.italic
                c = color_str(f.color)
                if c:
                    sd["color"] = c
        except Exception:
            pass
        try:
            pf = style.paragraph_format
            if pf:
                if pf.alignment is not None:
                    sd["alignment"] = str(pf.alignment)
                if pf.space_before is not None:
                    sd["space_before_pt"] = pf.space_before.pt
                if pf.space_after is not None:
                    sd["space_after_pt"] = pf.space_after.pt
                if pf.line_spacing is not None:
                    sd["line_spacing"] = pf.line_spacing
                if pf.line_spacing_rule is not None:
                    sd["line_spacing_rule"] = str(pf.line_spacing_rule)
                if pf.left_indent is not None:
                    sd["left_indent_in"] = emu_to_inches(pf.left_indent)
                if pf.keep_with_next is not None:
                    sd["keep_with_next"] = pf.keep_with_next
        except Exception:
            pass
        # Check for borders in style
        try:
            se = style.element
            pp = se.find(qn("w:pPr"))
            if pp is not None:
                pb = pp.find(qn("w:pBdr"))
                if pb is not None:
                    borders = {}
                    for side in ["top", "bottom", "left", "right"]:
                        b = pb.find(qn(f"w:{side}"))
                        if b is not None:
                            borders[side] = {
                                "val": b.get(qn("w:val")),
                                "sz": b.get(qn("w:sz")),
                                "color": b.get(qn("w:color")),
                            }
                    if borders:
                        sd["borders"] = borders
        except Exception:
            pass
        style_defs[style.name] = sd
    result["style_definitions"] = style_defs

    # Paragraphs
    paragraphs = []
    for idx, para in enumerate(doc.paragraphs):
        pdata = {
            "index": idx,
            "style": para.style.name if para.style else None,
            "text_preview": para.text[:100] if para.text else "",
        }
        pf = para.paragraph_format
        if pf.alignment is not None:
            pdata["alignment"] = str(pf.alignment)
        if pf.space_before is not None:
            pdata["space_before_pt"] = pf.space_before.pt
        if pf.space_after is not None:
            pdata["space_after_pt"] = pf.space_after.pt
        if pf.line_spacing is not None:
            pdata["line_spacing"] = pf.line_spacing
        if pf.left_indent is not None:
            pdata["left_indent_in"] = emu_to_inches(pf.left_indent)

        # Numbering
        pp = para._element.find(qn("w:pPr"))
        if pp is not None:
            np2 = pp.find(qn("w:numPr"))
            if np2 is not None:
                il = np2.find(qn("w:ilvl"))
                ni = np2.find(qn("w:numId"))
                pdata["num_level"] = il.get(qn("w:val")) if il is not None else None
                pdata["num_id"] = ni.get(qn("w:val")) if ni is not None else None
            # Borders
            pb = pp.find(qn("w:pBdr"))
            if pb is not None:
                borders = {}
                for side in ["top", "bottom", "left", "right"]:
                    b = pb.find(qn(f"w:{side}"))
                    if b is not None:
                        borders[side] = {
                            "val": b.get(qn("w:val")),
                            "sz": b.get(qn("w:sz")),
                            "color": b.get(qn("w:color")),
                        }
                if borders:
                    pdata["borders"] = borders

        # Runs
        runs = []
        for ri, run in enumerate(para.runs):
            rd = {"text_preview": run.text[:60]}
            f = run.font
            if f.name:
                rd["font_name"] = f.name
            if f.size:
                rd["font_size_pt"] = f.size.pt
            if f.bold is not None:
                rd["bold"] = f.bold
            if f.italic is not None:
                rd["italic"] = f.italic
            if f.underline is not None:
                rd["underline"] = f.underline
            if f.all_caps is not None:
                rd["all_caps"] = f.all_caps
            c = color_str(f.color)
            if c:
                rd["color"] = c
            runs.append(rd)
        pdata["runs"] = runs
        paragraphs.append(pdata)
    result["paragraphs"] = paragraphs

    # Numbering definitions
    numbering = []
    try:
        numbering_part = doc.part.numbering_part
        if numbering_part is not None:
            ne = numbering_part.element
            for an in ne.findall(qn("w:abstractNum")):
                aid = an.get(qn("w:abstractNumId"))
                levels = []
                for lvl in an.findall(qn("w:lvl")):
                    il = lvl.get(qn("w:ilvl"))
                    nf = lvl.find(qn("w:numFmt"))
                    lt = lvl.find(qn("w:lvlText"))
                    ld = {
                        "level": il,
                        "numFmt": nf.get(qn("w:val")) if nf is not None else None,
                        "lvlText": lt.get(qn("w:val")) if lt is not None else None,
                    }
                    pp = lvl.find(qn("w:pPr"))
                    if pp is not None:
                        ind = pp.find(qn("w:ind"))
                        if ind is not None:
                            left = ind.get(qn("w:left"))
                            hang = ind.get(qn("w:hanging"))
                            if left:
                                ld["indent_left_in"] = round(int(left) / 1440, 3)
                            if hang:
                                ld["hanging_in"] = round(int(hang) / 1440, 3)
                    rp = lvl.find(qn("w:rPr"))
                    if rp is not None:
                        rf = rp.find(qn("w:rFonts"))
                        if rf is not None:
                            ld["bullet_font"] = rf.get(qn("w:ascii"))
                    levels.append(ld)
                numbering.append({"abstractNumId": aid, "levels": levels})
    except Exception:
        pass
    result["numbering"] = numbering

    return result


def print_human_readable(data):
    print("=== PAGE SETUP ===")
    for s in data["page_setup"]:
        print(f"  Section {s['index']}: {s['page_width_in']}x{s['page_height_in']}in "
              f"Margins T={s['margin_top_in']} B={s['margin_bottom_in']} "
              f"L={s['margin_left_in']} R={s['margin_right_in']}")

    print("\n=== DOC DEFAULTS ===")
    for k, v in data["doc_defaults"].items():
        print(f"  {k}: {v}")

    print(f"\n=== PARA STYLES USED ({len(data['para_styles_used'])}) ===")
    for s in data["para_styles_used"]:
        print(f"  {s}")

    print("\n=== STYLE DEFINITIONS ===")
    for name, sd in data["style_definitions"].items():
        print(f"  [{name}] {sd}")

    print(f"\n=== PARAGRAPHS ({len(data['paragraphs'])}) ===")
    for p in data["paragraphs"]:
        extras = []
        for k in ["alignment", "space_before_pt", "space_after_pt", "line_spacing",
                   "left_indent_in", "num_level", "num_id", "borders"]:
            if k in p:
                extras.append(f"{k}={p[k]}")
        ext = " | ".join(extras) if extras else ""
        print(f"  P{p['index']} [{p['style']}] {ext}")
        print(f"    \"{p['text_preview'][:80]}\"")
        for r in p["runs"]:
            rd = {k: v for k, v in r.items() if k != "text_preview"}
            print(f"      [{rd}] \"{r['text_preview'][:50]}\"")

    print(f"\n=== NUMBERING ({len(data['numbering'])}) ===")
    for n in data["numbering"]:
        print(f"  AbstractNum {n['abstractNumId']}:")
        for l in n["levels"]:
            print(f"    {l}")


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python analyze_formatting.py <path.docx> [--json]")
        sys.exit(1)

    path = sys.argv[1]
    use_json = "--json" in sys.argv

    data = analyze(path)

    if use_json:
        print(json.dumps(data, indent=2, default=str))
    else:
        print_human_readable(data)
