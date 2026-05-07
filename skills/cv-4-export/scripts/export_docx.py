#!/usr/bin/env python3
# /// script
# requires-python = ">=3.10"
# dependencies = ["python-docx>=1.1.0"]
# ///
"""Generate a professionally formatted DOCX resume from structured JSON data.

Usage:
  uv run export_docx.py <cv_data.json> [--output <filename.docx>]
  python export_docx.py <cv_data.json> [--output <filename.docx>]

The JSON input schema is defined in skills/cv-4-export/SKILL.md.
Formatting matches the Resume_generic.docx reference document.
"""

import argparse
import json
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


# ---------------------------------------------------------------------------
# Formatting constants (extracted from Resume_generic.docx)
# ---------------------------------------------------------------------------
FONT_FAMILY = "Calibri"

# Page: A4 with narrow margins
PAGE_WIDTH_CM = 21.0
PAGE_HEIGHT_CM = 29.7
MARGIN_TOP_CM = 1.27        # 0.5 in
MARGIN_BOTTOM_CM = 1.27     # 0.5 in
MARGIN_LEFT_CM = 1.693      # 0.667 in
MARGIN_RIGHT_CM = 1.693     # 0.667 in

# Font sizes
SIZE_NAME = Pt(20)
SIZE_CONTACT = Pt(8.5)
SIZE_SECTION_HEADING = Pt(10)   # inherits from Normal default
SIZE_SUMMARY_BODY = Pt(9)
SIZE_TECH_SKILL = Pt(8.5)
SIZE_JOB_TITLE = Pt(9)
SIZE_COMPANY_LINE = Pt(8)
SIZE_BULLET = Pt(8.5)
SIZE_EDUCATION_TITLE = Pt(8.5)
SIZE_EDUCATION_SUBTITLE = Pt(8)
SIZE_EDUCATION_DETAIL = Pt(8)
SIZE_CERT = Pt(8)

# Colors
COLOR_NAME = RGBColor(0x1A, 0x1A, 0x1A)
COLOR_CONTACT = RGBColor(0x55, 0x55, 0x55)
COLOR_HEADING = RGBColor(0x24, 0x29, 0x2F)
COLOR_JOB_TITLE = RGBColor(0x1A, 0x1A, 0x1A)
COLOR_COMPANY = RGBColor(0x88, 0x88, 0x88)
COLOR_BLACK = RGBColor(0x00, 0x00, 0x00)

# Spacing (in points)
SP_NAME_AFTER = Pt(3)
SP_CONTACT_AFTER = Pt(1)
SP_LANGUAGES_AFTER = Pt(4)
SP_HEADING_BEFORE = Pt(10)
SP_HEADING_AFTER = Pt(3)
SP_SUMMARY_AFTER = Pt(2)
SP_TECH_AFTER = Pt(1)
SP_JOB_TITLE_BEFORE = Pt(5)
SP_JOB_TITLE_AFTER = Pt(1)
SP_COMPANY_AFTER = Pt(2)
SP_BULLET_AFTER = Pt(2)
SP_EDU_TITLE_BEFORE = Pt(4)
SP_EDU_TITLE_AFTER = Pt(1)
SP_EDU_SUB_AFTER = Pt(1)
SP_CERT_AFTER = Pt(1)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def set_run(run, font_size=None, bold=None, italic=None, color=None):
    """Apply formatting to a run."""
    run.font.name = FONT_FAMILY
    # Set east-asian font to Calibri as well for consistency
    r = run._element
    rPr = r.find(qn("w:rPr"))
    if rPr is None:
        rPr = OxmlElement("w:rPr")
        r.insert(0, rPr)
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), FONT_FAMILY)
    rFonts.set(qn("w:cs"), FONT_FAMILY)

    if font_size is not None:
        run.font.size = font_size
    if bold is not None:
        run.font.bold = bold
    if italic is not None:
        run.font.italic = italic
    if color is not None:
        run.font.color.rgb = color


def add_paragraph(doc, text="", font_size=None, bold=None, color=None,
                  space_before=None, space_after=None,
                  alignment=None, keep_with_next=False):
    """Add a paragraph with consistent formatting."""
    para = doc.add_paragraph()
    pf = para.paragraph_format
    if space_before is not None:
        pf.space_before = space_before
    if space_after is not None:
        pf.space_after = space_after
    if alignment is not None:
        pf.alignment = alignment
    if keep_with_next:
        pf.keep_with_next = True
    pf.line_spacing = 1.0

    if text:
        run = para.add_run(text)
        set_run(run, font_size=font_size, bold=bold, color=color)

    return para


def add_bottom_border(para, color_hex="C8C8C8", size="1"):
    """Add a thin bottom border to a paragraph (section heading style)."""
    pPr = para._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:color"), color_hex)
    bottom.set(qn("w:space"), "1")
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_bullet_paragraph(doc, text, font_size=SIZE_BULLET, space_after=SP_BULLET_AFTER):
    """Add a bullet point paragraph using Calibri bullet character and manual indent."""
    para = doc.add_paragraph()
    pf = para.paragraph_format
    pf.space_after = space_after
    pf.space_before = Pt(0)
    pf.line_spacing = 1.0

    # Set indentation: left=0.25in (0.635cm), hanging=0.167in (0.424cm)
    pf.left_indent = Cm(0.635)
    pf.first_line_indent = Cm(-0.424)

    # Bullet character
    bullet_run = para.add_run("\u2022\t")
    set_run(bullet_run, font_size=font_size)

    # Main text
    text_run = para.add_run(text)
    set_run(text_run, font_size=font_size)

    # Set tab stop at indent position for bullet alignment
    pPr = para._element.get_or_add_pPr()
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "left")
    tab.set(qn("w:pos"), str(int(0.635 * 567)))  # convert cm to twips
    tabs.append(tab)
    pPr.append(tabs)

    return para


def add_section_heading(doc, text):
    """Add an ALL-CAPS section heading with bottom border."""
    para = add_paragraph(
        doc, text.upper(),
        bold=True, color=COLOR_HEADING,
        space_before=SP_HEADING_BEFORE, space_after=SP_HEADING_AFTER,
        keep_with_next=True,
    )
    add_bottom_border(para)
    return para


# ---------------------------------------------------------------------------
# Section renderers
# ---------------------------------------------------------------------------

def render_header(doc, header):
    """Render the name, contact line, and languages."""
    # Name
    add_paragraph(doc, header["name"],
                  font_size=SIZE_NAME, bold=True, color=COLOR_NAME,
                  space_after=SP_NAME_AFTER)

    # Contact line
    add_paragraph(doc, header.get("contact_line", ""),
                  font_size=SIZE_CONTACT, color=COLOR_CONTACT,
                  space_after=SP_CONTACT_AFTER)

    # Languages (optional)
    if header.get("languages"):
        add_paragraph(doc, header["languages"],
                      font_size=SIZE_CONTACT, color=COLOR_CONTACT,
                      space_after=SP_LANGUAGES_AFTER)


def render_summary(doc, section):
    """Render the professional summary section."""
    add_section_heading(doc, section["heading"])
    add_paragraph(doc, section["body"],
                  font_size=SIZE_SUMMARY_BODY,
                  space_after=SP_SUMMARY_AFTER,
                  alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)


def render_skills_list(doc, section):
    """Render a key skills section as a compact paragraph."""
    add_section_heading(doc, section["heading"])
    items_text = ", ".join(section.get("items", []))
    add_paragraph(doc, items_text,
                  font_size=SIZE_TECH_SKILL,
                  space_after=SP_TECH_AFTER)


def render_technical_skills(doc, section):
    """Render categorized technical skills."""
    add_section_heading(doc, section["heading"])
    for cat in section.get("categories", []):
        para = add_paragraph(doc, space_after=SP_TECH_AFTER)
        label_run = para.add_run(cat["label"] + ": ")
        set_run(label_run, font_size=SIZE_TECH_SKILL, bold=True)
        value_run = para.add_run(cat["value"])
        set_run(value_run, font_size=SIZE_TECH_SKILL)


def render_experience(doc, section):
    """Render professional experience with roles and bullets."""
    add_section_heading(doc, section["heading"])
    for role in section.get("roles", []):
        # Job title
        add_paragraph(doc, role["title"],
                      font_size=SIZE_JOB_TITLE, bold=True, color=COLOR_JOB_TITLE,
                      space_before=SP_JOB_TITLE_BEFORE, space_after=SP_JOB_TITLE_AFTER,
                      keep_with_next=True)
        # Company line
        add_paragraph(doc, role["company_line"],
                      font_size=SIZE_COMPANY_LINE, color=COLOR_COMPANY,
                      space_after=SP_COMPANY_AFTER,
                      keep_with_next=True)
        # Bullets
        for bullet in role.get("bullets", []):
            add_bullet_paragraph(doc, bullet)


def render_education(doc, section):
    """Render education entries."""
    add_section_heading(doc, section["heading"])
    for i, entry in enumerate(section.get("entries", [])):
        sp_before = SP_EDU_TITLE_BEFORE if i > 0 else Pt(0)
        # Title (degree)
        add_paragraph(doc, entry["title"],
                      font_size=SIZE_EDUCATION_TITLE, bold=True, color=COLOR_JOB_TITLE,
                      space_before=sp_before, space_after=SP_EDU_TITLE_AFTER)
        # Subtitle (institution)
        add_paragraph(doc, entry.get("subtitle", ""),
                      font_size=SIZE_EDUCATION_SUBTITLE, color=COLOR_COMPANY,
                      space_after=SP_EDU_SUB_AFTER)
        # Detail (thesis, achievements)
        if entry.get("detail"):
            add_paragraph(doc, entry["detail"],
                          font_size=SIZE_EDUCATION_DETAIL,
                          space_after=SP_EDU_SUB_AFTER)


def render_certifications(doc, section):
    """Render certifications or awards."""
    add_section_heading(doc, section["heading"])
    for item in section.get("items", []):
        para = add_paragraph(doc, space_after=SP_CERT_AFTER)
        name_run = para.add_run(item["name"])
        set_run(name_run, font_size=SIZE_CERT, bold=True)
        if item.get("detail"):
            detail_run = para.add_run(item["detail"])
            set_run(detail_run, font_size=SIZE_CERT)


def render_bullets_section(doc, section):
    """Render a section that is just a heading + bullet list."""
    add_section_heading(doc, section["heading"])
    for bullet in section.get("bullets", []):
        add_bullet_paragraph(doc, bullet)


# ---------------------------------------------------------------------------
# Section type dispatch
# ---------------------------------------------------------------------------

SECTION_RENDERERS = {
    "summary": render_summary,
    "skills_list": render_skills_list,
    "technical_skills": render_technical_skills,
    "experience": render_experience,
    "education": render_education,
    "certifications": render_certifications,
    "awards": render_certifications,  # same format
    "bullets_section": render_bullets_section,
}


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def generate_docx(cv_data, output_path):
    """Generate a DOCX file from structured CV data."""
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Cm(PAGE_WIDTH_CM)
    section.page_height = Cm(PAGE_HEIGHT_CM)
    section.top_margin = Cm(MARGIN_TOP_CM)
    section.bottom_margin = Cm(MARGIN_BOTTOM_CM)
    section.left_margin = Cm(MARGIN_LEFT_CM)
    section.right_margin = Cm(MARGIN_RIGHT_CM)

    # Set docDefaults-level font (w:docDefaults/w:rPrDefault/w:rPr)
    styles_elem = doc.styles.element
    doc_defaults = styles_elem.find(qn("w:docDefaults"))
    if doc_defaults is None:
        doc_defaults = OxmlElement("w:docDefaults")
        styles_elem.insert(0, doc_defaults)
    rpr_default = doc_defaults.find(qn("w:rPrDefault"))
    if rpr_default is None:
        rpr_default = OxmlElement("w:rPrDefault")
        doc_defaults.append(rpr_default)
    dd_rPr = rpr_default.find(qn("w:rPr"))
    if dd_rPr is None:
        dd_rPr = OxmlElement("w:rPr")
        rpr_default.append(dd_rPr)
    dd_rFonts = dd_rPr.find(qn("w:rFonts"))
    if dd_rFonts is None:
        dd_rFonts = OxmlElement("w:rFonts")
        dd_rPr.insert(0, dd_rFonts)
    dd_rFonts.set(qn("w:ascii"), FONT_FAMILY)
    dd_rFonts.set(qn("w:hAnsi"), FONT_FAMILY)
    dd_rFonts.set(qn("w:eastAsia"), FONT_FAMILY)
    dd_rFonts.set(qn("w:cs"), FONT_FAMILY)

    # Set Normal style font
    style = doc.styles["Normal"]
    font = style.font
    font.name = FONT_FAMILY
    font.size = Pt(10)
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), FONT_FAMILY)
    rFonts.set(qn("w:hAnsi"), FONT_FAMILY)
    rFonts.set(qn("w:eastAsia"), FONT_FAMILY)
    rFonts.set(qn("w:cs"), FONT_FAMILY)

    # Set default paragraph spacing
    pf = style.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = 1.0

    # Remove the empty default paragraph that Document() creates
    if doc.paragraphs and not doc.paragraphs[0].text:
        p_element = doc.paragraphs[0]._element
        p_element.getparent().remove(p_element)

    # Render header
    render_header(doc, cv_data["header"])

    # Render sections
    for sec in cv_data.get("sections", []):
        renderer = SECTION_RENDERERS.get(sec["type"])
        if renderer:
            renderer(doc, sec)
        else:
            print(f"Warning: Unknown section type '{sec['type']}', skipping.",
                  file=sys.stderr)

    doc.save(str(output_path))
    print(f"DOCX generated: {output_path}")


def main():
    parser = argparse.ArgumentParser(description="Generate DOCX resume from JSON")
    parser.add_argument("input", help="Path to CV data JSON file")
    parser.add_argument("--output", "-o", default="resume_output.docx",
                        help="Output DOCX filename (default: resume_output.docx)")
    args = parser.parse_args()

    input_path = Path(args.input)
    if not input_path.exists():
        print(f"Error: Input file not found: {input_path}", file=sys.stderr)
        sys.exit(1)

    with open(input_path, "r", encoding="utf-8") as f:
        cv_data = json.load(f)

    generate_docx(cv_data, Path(args.output))


if __name__ == "__main__":
    main()
